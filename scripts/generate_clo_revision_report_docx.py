"""Generate the formal CLO revision report as a deterministic Word document."""

from __future__ import annotations

import json
import os
import re
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
AUDIT_PATH = ROOT / "data" / "clo_revision_audit_term_251_to_261.json"
CURRICULUM_PATH = ROOT / "data" / "ee_curriculum.json"
OUTPUT_PATH = ROOT / "output" / "CLO_Revision_Report_Term_251_to_261.docx"
TEMPLATE_PATH = ROOT / "templates" / "CLO-Revision-Report-Template.docx"

# Standard business brief, with the requested A4 portrait override.
FONT = "Arial"
NAVY = "17365D"
BLUE = "2E74B5"
INK = "202A35"
MUTED = "59636E"
GRID = "A9B3BE"
HEADER_FILL = "E8EDF3"
TABLE_WIDTH_DXA = 9400
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 90, "start": 120, "bottom": 90, "end": 120}
FIXED_TIME = datetime(2026, 1, 1, tzinfo=timezone.utc)

LABELS = {
    "unchanged": "Unchanged",
    "modified": "Modified",
    "renumbered": "Renumbered",
    "added": "Added",
    "omitted": "Omitted",
    "merged": "Merged",
    "split": "Split",
    "ambiguous": "Review required",
}


def set_run_font(run, *, size=None, bold=None, italic=None, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 11, 8, 4, NAVY),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGINS_DXA.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths):
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA} DXA: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "5")
        border.set(qn("w:color"), GRID)


def style_table(table, widths, *, header=True, font_size=8.5, center_columns=()):
    set_table_geometry(table, widths)
    set_table_borders(table)
    if header:
        set_repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        if header and row_index == 0:
            for cell in row.cells:
                set_cell_shading(cell, HEADER_FILL)
        for col_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index in center_columns else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=font_size, bold=(header and row_index == 0))


def add_text(cell, text, *, bold=False, italic=False, size=8.2, color=INK):
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph


def clear_cell(cell):
    paragraph = cell.paragraphs[0]
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def current_mapping(clo, pi_owners):
    domain = domain_name(clo.get("nqf_domain"))
    sos = [str(value) for value in clo.get("mapped_sos", [])]
    grouped = {so: [] for so in sos}
    for pi in clo.get("pi_codes", []):
        owner = pi_owners.get(pi)
        if owner not in grouped:
            raise ValueError(f"Invalid SO/PI relationship for {clo.get('current_clo_id')}: {pi}")
        grouped[owner].append(pi)
    values = [f"{so} ({', '.join(grouped[so])})" if grouped[so] else so for so in sos]
    return f"{domain} · {', '.join(values)}" if values else domain


def baseline_mapping(clo):
    domain = domain_name(clo.get("nqf_domain"))
    source = str(clo.get("source_so", ""))
    sos = []
    for value in re.findall(r"\d+", source):
        so = f"SO{value}"
        if so not in sos:
            sos.append(so)
    return f"{domain} · {', '.join(sos)}" if sos else domain


def domain_name(value):
    value = str(value or "")
    return "Knowledge" if value.startswith("Knowledge") else "Skills" if value.startswith("Skills") else "Values"


def add_clo_blocks(cell, items, side, pi_owners):
    clear_cell(cell)
    if not items:
        add_text(cell, "-", color=MUTED)
        return
    for index, item in enumerate(items):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(4 if index < len(items) - 1 else 1)
        clo_id = item["baseline_clo_ref"] if side == "baseline" else item["current_clo_id"]
        wording = item["clo_text"] if side == "baseline" else item["current_clo_text"]
        mapping = baseline_mapping(item) if side == "baseline" else current_mapping(item, pi_owners)
        run = paragraph.add_run(f"CLO {clo_id}\n")
        set_run_font(run, size=8.1, bold=True, color=NAVY)
        run = paragraph.add_run(f"{mapping}\n")
        set_run_font(run, size=7.7, italic=True, color=MUTED)
        run = paragraph.add_run(wording)
        set_run_font(run, size=8.1)


def comment_for(item, index):
    return item["comment"]


def add_heading(doc, text, level):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, *, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead, rest = text[: len(bold_lead)], text[len(bold_lead) :]
        set_run_font(paragraph.add_run(lead), size=10, bold=True)
        set_run_font(paragraph.add_run(rest), size=10)
    else:
        set_run_font(paragraph.add_run(text), size=10)
    return paragraph


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(35.56)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(17.78)
    section.right_margin = Mm(17.78)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(0)


def relationship_sort_key(item):
    current = item.get("current_clos", [])
    baseline = item.get("baseline_clos", [])
    clo_id = current[0].get("current_clo_id") if current else baseline[0].get("baseline_clo_ref") if baseline else ""
    parts = tuple(int(value) for value in re.findall(r"\d+", str(clo_id)))
    return parts, item.get("type") == "omitted"


def add_title_block(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(34)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(8)
    set_run_font(kicker.add_run("UNDERGRADUATE ELECTRICAL ENGINEERING PROGRAM"), size=10, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("CLO Revision Report"), size=24, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    set_run_font(subtitle.add_run("Term 251 ABET Submission to Term 261 Proposed Curriculum"), size=13, bold=True, color=MUTED)
    meta = doc.add_table(rows=3, cols=2)
    values = [
        ("Comparison", "Term 251 to Term 261"),
        ("Prepared by", "Curriculum Committee"),
        ("Review status", "Proposed for Approval"),
    ]
    for row, (label, value) in zip(meta.rows, values):
        clear_cell(row.cells[0]); clear_cell(row.cells[1])
        add_text(row.cells[0], label, bold=True, size=9, color=NAVY)
        add_text(row.cells[1], value, size=9)
    style_table(meta, [2200, 7200], header=False, font_size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def build_report(audit, curriculum):
    counts = audit["counts"]
    relationships = audit["relationships"]
    summaries = {item["course_code"]: item for item in audit["course_summaries"]}
    pi_owners = {pi: definition["so"] for pi, definition in curriculum["abet"]["performance_indicators"].items()}

    if audit["accounting"]["baseline_accounted"] != counts["baseline_clos"]:
        raise ValueError("Baseline CLO reconciliation is incomplete")
    if audit["accounting"]["current_accounted"] != counts["current_clos"]:
        raise ValueError("Current CLO reconciliation is incomplete")
    ee403 = [item for item in relationships if item["course_code"] == "EE 403"]
    if not any(item["type"] == "unchanged" and any(str(clo["current_clo_id"]) == "1.1" for clo in item["current_clos"]) for item in ee403):
        raise ValueError("EE 403 unchanged CLO 1.1 is missing from the structured audit")

    doc = Document(TEMPLATE_PATH)
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    configure_styles(doc)
    configure_page(doc)
    props = doc.core_properties
    props.title = "CLO Revision Report - Term 251 to Term 261"
    props.subject = "Formal CLO reconciliation report"
    props.author = "Curriculum Committee"
    props.created = FIXED_TIME
    props.modified = FIXED_TIME

    add_title_block(doc)
    add_heading(doc, "1. Purpose", 1)
    add_body(doc, "This report documents the revision of the CLOs of the Undergraduate Electrical Engineering Program from the Term 251 curriculum submitted to ABET to the proposed Term 261 curriculum. It provides a concise record for curriculum review and approval.")
    add_heading(doc, "2. Background", 1)
    add_body(doc, "The CLO review formed part of the program's continuous-improvement process initiated in December 2025 following ABET review and assessment discussions. Faculty review and refinement continued through Term 252 and subsequent stages, resulting in the proposed Term 261 CLO set.")
    add_heading(doc, "3. Rationale for CLO Review", 1)
    add_body(doc, "The review was undertaken to improve CLO clarity and measurability and to strengthen alignment with SO assessment and the program's assessment framework. The revisions support assessment improvement; they do not imply that ABET prescribed specific CLO wording.")
    add_heading(doc, "4. Scope", 1)
    add_body(doc, "The scope of this review is limited to the 16 Undergraduate EE courses represented in both the Term 251 and Term 261 curricula. Elective courses and non-EE College/supporting courses are excluded from the comparison.")
    scope = doc.add_table(rows=1, cols=2)
    scope.rows[0].cells[0].text = "Course"
    scope.rows[0].cells[1].text = "Title"
    for item in audit["course_summaries"]:
        row = scope.add_row().cells
        row[0].text = item["course_code"]
        row[1].text = item["course_title"]
    style_table(scope, [1800, 7600], center_columns=(0,))

    add_heading(doc, "5. Summary of CLO Changes", 1)
    summary = doc.add_table(rows=1, cols=2)
    summary.rows[0].cells[0].text = "Category"
    summary.rows[0].cells[1].text = "Count"
    summary_items = [
        ("Term 251 CLOs", "baseline_clos"), ("Term 261 CLOs", "current_clos"),
        ("Unchanged", "unchanged"), ("Modified", "modified"),
        ("Renumbered", "renumbered"), ("Added", "added"),
        ("Omitted", "omitted"), ("Merged", "merge_cases"),
        ("Split", "split_cases"), ("Ambiguous", "ambiguous"),
    ]
    for label, key in summary_items:
        row = summary.add_row().cells
        row[0].text = label
        row[1].text = str(counts[key])
    style_table(summary, [7000, 2400], center_columns=(1,))
    add_body(doc, f"The comparison accounts for all {counts['baseline_clos']} Term 251 CLOs and all {counts['current_clos']} Term 261 CLOs. Merged relationships are counted once and are not duplicated as omissions or additions.")

    doc.add_page_break()
    add_heading(doc, "6. Detailed CLO Comparison by Course", 1)
    for code in audit["courses_with_changes"]:
        heading = add_heading(doc, f"{code} - {summaries[code]['course_title']}", 2)
        heading.paragraph_format.page_break_before = False
        table = doc.add_table(rows=1, cols=4)
        for cell, text in zip(table.rows[0].cells, ("Change", "Term 251 CLO", "Term 261 CLO", "Brief Justification / Comment")):
            cell.text = text
        entries = sorted((item for item in relationships if item["course_code"] == code), key=relationship_sort_key)
        for index, item in enumerate(entries):
            row = table.add_row().cells
            row[0].text = LABELS[item["type"]]
            add_clo_blocks(row[1], item["baseline_clos"], "baseline", pi_owners)
            add_clo_blocks(row[2], item["current_clos"], "current", pi_owners)
            row[3].text = comment_for(item, index)
        style_table(table, [1400, 2650, 2650, 2700], font_size=8.1, center_columns=(0,))
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    add_heading(doc, "7. Courses with No CLO Changes", 1)
    unchanged = doc.add_table(rows=1, cols=2)
    unchanged.rows[0].cells[0].text = "Course Code"
    unchanged.rows[0].cells[1].text = "Course Title"
    for code in audit["courses_without_clo_changes"]:
        row = unchanged.add_row().cells
        row[0].text = code
        row[1].text = summaries[code]["course_title"]
    style_table(unchanged, [2200, 7200], center_columns=(0,))

    add_heading(doc, "8. Impact on ABET Assessment", 1)
    add_body(doc, "The revised framework strengthens the alignment and traceability among CLOs, SOs, PIs, and assessment evidence. This supports systematic continuous improvement while keeping the detailed comparison focused on CLO wording and structure.")
    doc.add_page_break()
    add_heading(doc, "9. Approval Status", 1)
    approval = doc.add_table(rows=1, cols=2)
    approval.rows[0].cells[0].text = "Item"
    approval.rows[0].cells[1].text = "Status / Approval"
    for label, value in (
        ("Program", "B.Sc. Electrical Engineering"),
        ("Curriculum Term", "261"),
        ("Prepared by", "Curriculum Committee"),
        ("Review Status", "Proposed for Approval"),
        ("College Curriculum Committee", ""),
        ("Approval Date", ""),
        ("Institutional Curriculum Committee, if required", ""),
        ("Approval Date", ""),
    ):
        row = approval.add_row().cells
        row[0].text = label
        row[1].text = value
    style_table(approval, [3800, 5600])
    add_heading(doc, "10. Conclusion", 1)
    add_body(doc, "This report compares the Term 251 CLO baseline submitted to ABET with the proposed Term 261 CLO framework. The changes reflect the program's continuous-improvement and faculty-review process and strengthen CLO clarity, measurability, and alignment with ABET assessment. The Term 261 CLO set is presented for the required curriculum approval.")
    return doc


def normalize_zip(path):
    """Normalize ZIP member order and timestamps for byte-stable output."""
    fd, temp_name = tempfile.mkstemp(suffix=".docx", dir=path.parent)
    os.close(fd)
    temp_path = Path(temp_name)
    try:
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as target:
            for name in sorted(source.namelist()):
                info = zipfile.ZipInfo(name, date_time=(2026, 1, 1, 0, 0, 0))
                info.compress_type = zipfile.ZIP_DEFLATED
                info.external_attr = 0o600 << 16
                target.writestr(info, source.read(name))
        temp_path.replace(path)
    finally:
        temp_path.unlink(missing_ok=True)


def main():
    audit = json.loads(AUDIT_PATH.read_text(encoding="utf-8"))
    curriculum = json.loads(CURRICULUM_PATH.read_text(encoding="utf-8"))
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_report(audit, curriculum)
    document.save(OUTPUT_PATH)
    normalize_zip(OUTPUT_PATH)
    print(json.dumps({"output": str(OUTPUT_PATH.relative_to(ROOT)), "source": str(AUDIT_PATH.relative_to(ROOT))}, indent=2))


if __name__ == "__main__":
    main()
